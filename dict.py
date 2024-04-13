import sys
import os
from pprint import pprint
from pypdf import PdfReader
import re
from termcolor import colored
# import aspose.words as aw
from docx import Document
import copy

sys.path.append(os.path.dirname(__file__) + '/CambridgeDict/cambridge_parser')
from CambridgeDict.cambridge_parser import define

def extract_pdf(pdf: str):
    reader = PdfReader(pdf+'.pdf') 
    print(len(reader.pages))  
    page = reader.pages[1]
    text = ''
    
    for page in reader.pages:
        text += page.extract_text()
    return text 

def del_dupl(x):
    return list(dict.fromkeys(x))

def create_alph(words):
    alphabet = []
    for ind in range(26): alphabet.append([])
    
    for word in words:
        if word[0].isalpha():
            ind = ord(word[0].lower()) - ord('a')
            alphabet[ind].append(word.lower()) if ind < 26 else words.remove(word)
        else:
            words.remove(word)
    return alphabet

def create_dict(text):
    words_list = re.findall(r"[\w']+", text)
    words_list.sort()
    words_list = del_dupl(words_list)
    
    for _ in range(10):
        for word in words_list:
            if word.isnumeric(): 
                words_list.remove(word)
    return words_list

# def generate_doc():
    # create document object
    doc = aw.Document()
    # create a document builder object
    builder = aw.DocumentBuilder(doc)
    # add text to the document
    builder.write("Hello world!")
    # save document

    builder.writeln("Item 1")
    builder.writeln("Item 2")

    # set indentation for next level of list
    builder.list_format.list_indent()
    builder.writeln("Item 2.1")
    builder.writeln("Item 2.2")

    builder.list_format.list_outdent()
    builder.writeln("Item 3")
    # remove numbers
    builder.list_format.remove_numbers()

    doc.save("out.docx")

L = "List Bullet"
LL = "List Bullet 2"
LL = "List Bullet 3"

def add_word(doc, word, expl, defin, text_exmpl, dict_exmpl):
    doc.add_paragraph(word + f' [{expl}]', style=L)
    doc.add_paragraph(' - ' + defin, style=LL)
    doc.add_paragraph(f'"{text_exmpl}"', style=LL)
    doc.add_paragraph(f'"{dict_exmpl}"', style=LL)
    doc.add_paragraph(f'" "', style=LL)
    return

def generate_doc():
    doc = Document()
    doc.add_heading('Personal Vocabular, MIPT 2024')
    doc.add_heading("a. Definition from an Eng-Eng dictionary (Macmillan / Cambridge / Collins / etc / or field-specific dictionary).\n"
                    "b. Sentence of usage from the article\n"
                    "c. Sentence of usage from the dictionary\n"  
                    "d. Your own example sentence\n", 2
                    )
    # doc.add_paragraph("It was a dark and stormy night.")
    # doc.add_paragraph()
    # <docx.text.paragraph.Paragraph object at 0x10f19e760>
    # doc.save("Personal Vocabular.docx")
    return doc

SPEC = {'e':['exit()', 'e()'], 'b':['back()', 'b()']}

# message assistant
def mess_ass(sentences=None, doc=None):
    assert(sentences)

    print(colored('Input the letter to show possible words:', 'green'))
    while(True):
        letter = input().lower()
        if letter in SPEC['e']:
            return False
        
        if len(letter) == 1 and ord(letter) - ord('a') < 26 and ord(letter) - ord('a') >= 0:  
            print(ord(letter) - ord('a'))
            pprint(alphabet[ord(letter) - ord('a')])
            break
        else:
            print('Incorrect input, try again.')
    
    pdf_sent = None
    word = None
    print(colored('Input the word to create its definition:', 'green'))
    while(True):
        word = input().lower()
        
        if word in SPEC['b']:
            return True
        if word not in alphabet[ord(word[0]) - ord('a')]:
            print('Incorrect word.')
            continue
        tokens = '.!?-,'
        exit_for = False
        for sent in sentences:
            if ' '+word in sent.lower():
                pdf_sent=sent
                break
            else:
                for t in tokens:
                    if ' '+word+t in sent.lower() or '-'+word+t in sent.lower(): 
                        exit_for=True
                        break
                if exit_for:
                    break
        break
    print(colored('success:', 'green'), pdf_sent)

    if pdf_sent is None:
        print(colored('Cannot find this word in pdf sentences :))). Use your fingers to find!', 'red'))

    res = define(word=word)
    pprint(res) 
    # for key in res[0][word]:
    #     definitions = key['data']['definitions'] 
    #     examples = key['data']['examples']
    #     print(key['POS'][0])
    #     print(definitions)
    #     print(examples)
    part, definitions, examples, num = None, None, None, 0
    cambridge_name = list(res[0].keys())[0]
    # print(cambridge_name[0])
    if len(res[0][cambridge_name]) > 1:
        print(colored('Choose the part, input example number from 1 to {}'.format(len(res[0][cambridge_name])), 'red'))
        pprint([key['POS'][0] for key in res[0][cambridge_name]])
        
        num = int(input())-1

    part = res[0][cambridge_name][num]['POS'][0]
    definitions = res[0][cambridge_name][num]['data']['definitions'] 
    examples = res[0][cambridge_name][num]['data']['examples']
        

    if not examples[0]:
        examples = None
    else:
        print(colored('Input example number from 1 to {}'.format(len(examples[0])), 'red'))
        pprint(examples[0])
        examples = examples[0][int(input())-1]

    add_word(doc, word, part, definitions[0], pdf_sent, examples)# key['POS'][0], key['data']['definitions'], key['data']['examples'][0])
    return True

if __name__ == '__main__':

    print(colored('Hello! Please, inter the pdf file name', 'green'), end='')
    print(colored('(without .pdf)', 'red'), end='')
    print(colored(':', 'green'))

    pdf_file = input()
    try: 
        pdf_text = extract_pdf(pdf_file) #('Thermal limits')
    except:
        print(colored('Invalid file (check name of file or its path)', 'red'))
    else:
    # if hasattr(pdf_text):
        text  = copy.copy(pdf_text)
        sentences = [key.replace("\n", "") for key in\
                                [x for x in re.split("[//.|//!|//?]", text) if x!=""]]

        words = create_dict(pdf_text)
            # print(len(words))
            # pprint(words)
        alphabet = create_alph(words)
        doc = generate_doc()

        print(colored('Shortcuts:', 'red'))
        print(colored('\t exit(), e() - stop to add words into DOC file and exit.\n'
                    '\t back(), b() - go back from words to select letter.\n',\
                    'light_blue'))
        while(mess_ass(sentences, doc)): continue
        
        doc.save('vocabulary.docx')

        # pprint(alphabet)

        # res = define(word=word)
        # pprint(res)
    # doc_name = ' '
    # generate_doc(doc_name=doc_name)



